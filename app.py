from __future__ import annotations

import io
import json
import hashlib
import re
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
import streamlit as st
import yfinance as yf
import feedparser

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    A4 = None
    canvas = None


APP_TITLE = "PM Workbench"
DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)

CACHE_FILE = DATA_DIR / "ai_cache.json"
USERS_FILE = DATA_DIR / "users.json"

DEFAULT_TICKERS = {
    "VNINDEX": "^VNINDEX",
    "S&P 500": "^GSPC",
    "Nasdaq": "^IXIC",
    "US 10Y Yield": "^TNX",
    "Gold": "GC=F",
    "Oil (WTI)": "CL=F",
    "USD Index": "DX-Y.NYB",
    "USD/VND": "VND=X",
}

VN_EXPERT_RECOMMENDATIONS = [
    {"ticker": "FPT", "action": "Buy", "source": "SSI"},
    {"ticker": "FPT", "action": "Buy", "source": "VNDirect"},
    {"ticker": "VCB", "action": "Buy", "source": "Fund"},
    {"ticker": "MBB", "action": "Buy", "source": "Fund"},
    {"ticker": "HPG", "action": "Watch", "source": "Research"},
    {"ticker": "SSI", "action": "Watch", "source": "Broker"},
    {"ticker": "MWG", "action": "Buy", "source": "Research"},
]

VN_SECTOR_MAP = {
    "FPT": "Technology",
    "VCB": "Banks",
    "MBB": "Banks",
    "HPG": "Steel",
    "SSI": "Securities",
    "MWG": "Consumer",
}

MODEL_PORTFOLIO = {
    "Core": [
        {"ticker": "FPT.VN", "weight": 0.20},
        {"ticker": "VCB.VN", "weight": 0.20},
        {"ticker": "MBB.VN", "weight": 0.15},
    ],
    "Tactical": [
        {"ticker": "MWG.VN", "weight": 0.10},
        {"ticker": "SSI.VN", "weight": 0.10},
        {"ticker": "HPG.VN", "weight": 0.10},
    ],
    "Hedge": [
        {"ticker": "Cash", "weight": 0.10},
        {"ticker": "USD", "weight": 0.05},
    ],
}

RSS_FEEDS = {
    "VnExpress Kinh doanh": "https://vnexpress.net/rss/kinh-doanh.rss",
    "VnExpress Chứng khoán": "https://vnexpress.net/rss/chung-khoan.rss",
    "Vietstock Chứng khoán": "https://vietstock.vn/rss/chung-khoan.rss",
    "Reuters Markets": "https://feeds.reuters.com/reuters/businessNews",
}

PROMPT_MASTER = """
You are a CIO / Portfolio Manager at an institutional fund.

Your job is NOT to summarize news.
Your job is to:
- interpret market information
- form an investment view
- translate it into portfolio actions

STRICT:
- Output must be strictly bilingual Vietnamese-English
- Every key point must appear in Vietnamese first, then English immediately after
- Never output English only
- No generic statements
- Always include:
  1. What changed
  2. Why it matters
  3. Portfolio implication
  4. Risk
  5. Trigger

STYLE:
- Institutional
- Concise
- Action-oriented
"""


def load_json(path: Path, default):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def save_json(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def ensure_files():
    if not USERS_FILE.exists():
        save_json(
            USERS_FILE,
            [{"username": "admin", "password_hash": hash_password("admin123"), "role": "admin"}],
        )
    if not CACHE_FILE.exists():
        save_json(CACHE_FILE, {})


def authenticate(username: str, password: str) -> Optional[dict]:
    users = load_json(USERS_FILE, [])
    pwd = hash_password(password)
    for u in users:
        if u["username"] == username and u["password_hash"] == pwd:
            return u
    return None


def get_secret(name: str, default: str = "") -> str:
    try:
        return str(st.secrets.get(name, default))
    except Exception:
        return default


def cached_ai_call(api_key: str, model: str, system_prompt: str, user_prompt: str, max_output_tokens: int = 1800) -> str:
    if not api_key or OpenAI is None:
        return ""
    cache = load_json(CACHE_FILE, {})
    key = hashlib.md5((model + system_prompt + user_prompt).encode("utf-8")).hexdigest()
    if key in cache:
        return cache[key]
    try:
        client = OpenAI(api_key=api_key)
        resp = client.responses.create(
            model=model,
            instructions=system_prompt,
            input=user_prompt,
            max_output_tokens=max_output_tokens,
        )
        out = getattr(resp, "output_text", "").strip()
        if out:
            cache[key] = out
            save_json(CACHE_FILE, cache)
        return out
    except Exception:
        return ""


def export_docx(title: str, content: str) -> bytes | None:
    if Document is None:
        return None
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def export_pdf(title: str, content: str) -> bytes | None:
    if canvas is None or A4 is None:
        return None
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    _, h = A4
    y = h - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, title[:90])
    y -= 24
    c.setFont("Helvetica", 10)
    for raw in content.split("\n"):
        chunks = [raw[i:i+110] for i in range(0, len(raw), 110)] or [""]
        for line in chunks:
            if y < 40:
                c.showPage()
                y = h - 40
                c.setFont("Helvetica", 10)
            c.drawString(40, y, line)
            y -= 14
    c.save()
    buf.seek(0)
    return buf.read()


def fetch_market_snapshot(tickers: Dict[str, str]) -> pd.DataFrame:
    rows = []
    for name, ticker in tickers.items():
        try:
            hist = yf.Ticker(ticker).history(period="2d", interval="1d", auto_adjust=False)
            if hist.empty or len(hist) < 2:
                continue
            last_close = float(hist.iloc[-1]["Close"])
            prev_close = float(hist.iloc[-2]["Close"])
            chg = round((last_close - prev_close) / prev_close * 100, 2) if prev_close else 0.0
            rows.append({"Asset": name, "Ticker": ticker, "Price": round(last_close, 2), "ChangePct": chg})
        except Exception:
            continue
    return pd.DataFrame(rows)


def fetch_last_price_for_ticker(ticker: str) -> Optional[float]:
    try:
        hist = yf.Ticker(ticker).history(period="5d", interval="1d", auto_adjust=False)
        if hist.empty:
            return None
        return float(hist.iloc[-1]["Close"])
    except Exception:
        return None


def fetch_news() -> pd.DataFrame:
    rows = []
    for source, url in RSS_FEEDS.items():
        try:
            feed = feedparser.parse(url)
            for entry in feed.entries[:6]:
                title = entry.get("title", "")
                summary = re.sub("<.*?>", "", entry.get("summary", "") or "")
                link = entry.get("link", "")
                txt = f"{title} {summary}".lower()
                region = "Vietnam" if any(x in txt or x in source.lower() for x in ["vnindex", "vietnam", "viet", "vnexpress", "vietstock"]) else "Global"
                asset = (
                    "Fixed Income" if any(x in txt for x in ["bond", "yield", "lãi suất", "fixed income"])
                    else "Commodity / FX" if any(x in txt for x in ["oil", "gold", "usd", "fx", "tỷ giá"])
                    else "Equity"
                )
                impact = 1 + sum(int(k in txt) for k in ["vnindex", "usd", "oil", "fpt", "vcb", "mbb", "hpg"])
                rows.append({
                    "Source": source,
                    "Title": title,
                    "Summary": summary,
                    "Link": link,
                    "Region": region,
                    "AssetClass": asset,
                    "VNImpact": min(impact, 5),
                })
        except Exception:
            continue
    return pd.DataFrame(rows)


def build_market_highlights(market_df: pd.DataFrame) -> str:
    if market_df.empty:
        return "No market data / Không có dữ liệu."
    return "\n".join([f"- {r.Asset}: {r.Price} ({r.ChangePct}%)" for _, r in market_df.iterrows()])


def build_news_brief(news_df: pd.DataFrame) -> str:
    if news_df.empty:
        return "No news / Không có tin."
    vn = news_df[news_df["Region"] == "Vietnam"].sort_values("VNImpact", ascending=False).head(4)
    gl = news_df[news_df["Region"] == "Global"].sort_values("VNImpact", ascending=False).head(4)
    lines = ["TIN VIỆT NAM / VIETNAM NEWS:"]
    lines += [f"- [{r.AssetClass}] {r.Title}" for _, r in vn.iterrows()] or ["- None"]
    lines += ["", "TIN QUỐC TẾ / GLOBAL NEWS:"]
    lines += [f"- [{r.AssetClass}] {r.Title}" for _, r in gl.iterrows()] or ["- None"]
    return "\n".join(lines)


def get_change(market_df: pd.DataFrame, name: str) -> float:
    row = market_df.loc[market_df["Asset"] == name]
    if row.empty:
        return 0.0
    try:
        return float(row.iloc[0]["ChangePct"])
    except Exception:
        return 0.0


def summarize_consensus() -> pd.DataFrame:
    df = pd.DataFrame(VN_EXPERT_RECOMMENDATIONS)
    out = []
    for ticker, grp in df.groupby("ticker"):
        buy = int((grp["action"].str.lower() == "buy").sum())
        watch = int((grp["action"].str.lower() == "watch").sum())
        mentions = len(grp)
        consensus = "Buy / Mua" if buy >= watch else "Watch / Theo dõi"
        out.append({
            "Ticker": ticker,
            "Mentions": mentions,
            "Buy": buy,
            "Watch": watch,
            "Consensus": consensus,
            "Sources": ", ".join(sorted(grp["source"].unique().tolist()))
        })
    return pd.DataFrame(out).sort_values(["Mentions", "Buy"], ascending=False).reset_index(drop=True)


def build_consensus_text(df: pd.DataFrame) -> str:
    if df.empty:
        return "No consensus / Chưa có consensus."
    lines = ["VIETNAM STOCK CONSENSUS / ĐỒNG THUẬN CỔ PHIẾU VIỆT NAM"]
    for _, r in df.head(8).iterrows():
        lines.append(f"- {r['Ticker']}: {r['Consensus']} | Mentions={r['Mentions']} | Sources={r['Sources']}")
    return "\n".join(lines)


def compute_allocation(market_df: pd.DataFrame, consensus_df: pd.DataFrame) -> dict:
    spx = get_change(market_df, "S&P 500")
    vn = get_change(market_df, "VNINDEX")
    usd = get_change(market_df, "USD Index")
    oil = get_change(market_df, "Oil (WTI)")
    y10 = get_change(market_df, "US 10Y Yield")

    regime_score = (1 if spx > 0.5 else 0) + (1 if vn > 0.7 else 0) - (1 if usd > 0.5 else 0) - (1 if oil > 1.5 else 0) - (1 if y10 > 0.5 else 0)
    regime = "Risk-on" if regime_score >= 1 else "Risk-off" if regime_score <= -1 else "Transition"

    alloc = {
        "Regime": regime,
        "Global Equity": {"View": "Overweight" if spx > 0.5 and y10 <= 0.3 else "Neutral", "Conviction": "Medium", "PortfolioFit": "Core risk asset"},
        "Vietnam Equity": {"View": "Overweight" if vn > 0.7 else "Neutral", "Conviction": "High" if consensus_df["Buy"].sum() >= 4 else "Medium", "PortfolioFit": "Core + tactical alpha"},
        "Global Fixed Income": {"View": "Underweight" if y10 > 0.5 or oil > 1.5 else "Neutral", "Conviction": "Medium", "PortfolioFit": "Hedge / duration"},
        "Gold": {"View": "Overweight" if regime == "Risk-off" or get_change(market_df, "Gold") > 0.5 else "Neutral", "Conviction": "Medium", "PortfolioFit": "Defensive hedge"},
        "USD": {"View": "Long" if usd > 0.5 else "Neutral", "Conviction": "Medium", "PortfolioFit": "Macro hedge"},
    }
    return alloc


def format_allocation(alloc: dict) -> str:
    lines = ["PORTFOLIO ALLOCATION / PHÂN BỔ DANH MỤC", f"- Regime / Chế độ: {alloc['Regime']}"]
    for k, v in alloc.items():
        if k == "Regime":
            continue
        lines += [
            "",
            f"{k}:",
            f"- View: {v['View']}",
            f"- Conviction: {v['Conviction']}",
            f"- Portfolio fit: {v['PortfolioFit']}",
        ]
    return "\n".join(lines)


def estimate_target_zone(asset: str, market_df: pd.DataFrame) -> str:
    if asset == "VNINDEX":
        row = market_df.loc[market_df["Asset"] == "VNINDEX"]
        if not row.empty:
            px = float(row.iloc[0]["Price"])
            return f"{round(px*1.02,1)} - {round(px*1.05,1)}"
    if asset.endswith(".VN"):
        px = fetch_last_price_for_ticker(asset)
        if px:
            return f"{round(px*1.05,2)} - {round(px*1.12,2)}"
    if asset == "Oil":
        return "+4% đến +8%"
    if asset == "Gold":
        return "+3% đến +6%"
    if asset == "USD":
        return "DXY +1% đến +3%"
    return "Tactical upside zone"


def score_idea(asset: str, market_df: pd.DataFrame, consensus_df: pd.DataFrame) -> dict:
    score = 5.0
    why = []

    if asset == "Oil" and get_change(market_df, "Oil (WTI)") > 1:
        score += 2
        why.append("Oil momentum strong / Động lượng dầu mạnh")
    if asset == "Gold" and get_change(market_df, "Gold") > 0.5:
        score += 1
        why.append("Defensive demand / Nhu cầu phòng thủ")
    if asset == "USD" and get_change(market_df, "USD Index") > 0.5:
        score += 2
        why.append("USD strength / USD mạnh")
    if asset == "VNINDEX" and get_change(market_df, "VNINDEX") > 0.7:
        score += 1.5
        why.append("VN momentum / Động lượng VN")
    if asset.endswith(".VN") and get_change(market_df, "VNINDEX") > 0.7:
        score += 1.0

    ticker_plain = asset.replace(".VN", "")
    if not consensus_df.empty and ticker_plain in consensus_df["Ticker"].tolist():
        score += 0.5
        why.append("Expert consensus / Đồng thuận chuyên gia")

    conviction = "High" if score >= 8 else "Medium" if score >= 6 else "Low"
    action = "Buy" if score >= 8 else "Add" if score >= 7 else "Hold" if score >= 5 else "Trim" if score >= 4 else "Exit"
    target = estimate_target_zone(asset, market_df)
    fit = "Hedge" if asset in ["Gold", "USD"] else "Core + tactical" if asset in ["VNINDEX", "Global Equity"] else "Alpha / stock selection"

    return {
        "Asset": asset,
        "Action": action,
        "Score": round(min(score, 10), 1),
        "Conviction": conviction,
        "WhyNow": "; ".join(why) if why else "Neutral setup / Thiết lập trung tính",
        "PortfolioFit": fit,
        "TargetZone": target,
    }


def build_trade_ideas(market_df: pd.DataFrame, consensus_df: pd.DataFrame) -> pd.DataFrame:
    universe = ["Oil", "Gold", "USD", "VNINDEX", "FPT.VN", "VCB.VN", "MBB.VN", "HPG.VN", "SSI.VN", "MWG.VN", "Global Equity"]
    rows = [score_idea(x, market_df, consensus_df) for x in universe]
    return pd.DataFrame(rows).sort_values("Score", ascending=False).reset_index(drop=True)


def format_trade_ideas(df: pd.DataFrame) -> str:
    if df.empty:
        return "No trade ideas / Chưa có trade ideas."
    lines = ["TRADE IDEAS / Ý TƯỞNG ĐẦU TƯ"]
    for i, r in df.head(8).iterrows():
        lines += [
            "",
            f"#{i+1} {r['Asset']}",
            f"- Action / Hành động: {r['Action']}",
            f"- Score: {r['Score']}",
            f"- Conviction: {r['Conviction']}",
            f"- Why now / Vì sao lúc này: {r['WhyNow']}",
            f"- Portfolio fit / Vai trò danh mục: {r['PortfolioFit']}",
            f"- Target zone / Vùng mục tiêu: {r['TargetZone']}",
        ]
    return "\n".join(lines)


def flatten_model_portfolio() -> List[dict]:
    rows = []
    for bucket, items in MODEL_PORTFOLIO.items():
        for x in items:
            rows.append({"Bucket": bucket, **x})
    return rows


def fetch_return_pct_for_item(ticker: str, market_df: pd.DataFrame) -> float:
    if ticker == "Cash":
        return 0.0
    if ticker == "USD":
        return get_change(market_df, "USD Index")
    df = fetch_market_snapshot({ticker: ticker})
    if df.empty:
        return 0.0
    return float(df.iloc[0]["ChangePct"])


def compute_benchmark(market_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    total = 0.0
    for item in flatten_model_portfolio():
        ret = fetch_return_pct_for_item(item["ticker"], market_df)
        contrib = item["weight"] * ret
        total += contrib
        rows.append({
            "Bucket": item["Bucket"],
            "Ticker": item["ticker"],
            "Weight": item["weight"],
            "ReturnPct": round(ret, 2),
            "Contribution": round(contrib, 2),
        })
    out = pd.DataFrame(rows)
    vn = get_change(market_df, "VNINDEX")
    out.attrs["portfolio_return"] = round(total, 2)
    out.attrs["vnindex_return"] = round(vn, 2)
    out.attrs["excess_return"] = round(total - vn, 2)
    return out


def build_heatmap(trade_df: pd.DataFrame) -> pd.DataFrame:
    if trade_df.empty:
        return pd.DataFrame()
    df = trade_df.copy()
    df["Ticker"] = df["Asset"].str.replace(".VN", "", regex=False)
    df["Sector"] = df["Ticker"].map(VN_SECTOR_MAP).fillna("Other")
    df["ConvictionNum"] = df["Conviction"].map({"High": 3, "Medium": 2, "Low": 1}).fillna(1)
    out = df.groupby("Sector").agg(
        IdeaCount=("Ticker", "count"),
        AvgScore=("Score", "mean"),
        AvgConviction=("ConvictionNum", "mean")
    ).reset_index()
    out["Heat"] = out["AvgConviction"].apply(lambda x: "Very High" if x >= 2.7 else "High" if x >= 2.2 else "Medium" if x >= 1.6 else "Low")
    out["AvgScore"] = out["AvgScore"].round(2)
    out["AvgConviction"] = out["AvgConviction"].round(2)
    return out.sort_values(["AvgConviction", "AvgScore"], ascending=False)


def generate_action_signals(trade_df: pd.DataFrame) -> pd.DataFrame:
    if trade_df.empty:
        return pd.DataFrame()
    model_tickers = [x["ticker"].upper() for x in flatten_model_portfolio()]
    rows = []
    for _, r in trade_df.iterrows():
        ticker = str(r["Asset"]).replace(".VN", "")
        in_portfolio = (ticker + ".VN").upper() in model_tickers
        action = r["Action"]
        if action == "Buy" and not in_portfolio:
            reason = "High conviction, not yet in portfolio"
        elif action in ["Buy", "Add"] and in_portfolio:
            action = "Add"
            reason = "Already held, setup still constructive"
        elif action == "Trim":
            reason = "Score softens"
        elif action == "Exit":
            reason = "Low score, weak setup"
        else:
            action = "Hold"
            reason = "No major change"
        rows.append({"Ticker": ticker, "Action": action, "Reason": reason, "Score": r["Score"], "Conviction": r["Conviction"]})
    return pd.DataFrame(rows)


def build_memo(allocation_text: str, trade_df: pd.DataFrame, action_df: pd.DataFrame, benchmark_df: pd.DataFrame) -> str:
    top_names = ", ".join(trade_df.head(3)["Asset"].tolist()) if not trade_df.empty else "N/A"
    actions = ", ".join([f"{r['Ticker']}:{r['Action']}" for _, r in action_df.head(4).iterrows()]) if not action_df.empty else "No action"
    return f"""PORTFOLIO MEMO / GHI NHỚ DANH MỤC

Regime:
{allocation_text.splitlines()[1] if allocation_text else "N/A"}

Benchmark:
- Model portfolio return: {benchmark_df.attrs.get('portfolio_return', 0.0)}%
- VNINDEX return: {benchmark_df.attrs.get('vnindex_return', 0.0)}%
- Excess return: {benchmark_df.attrs.get('excess_return', 0.0)}%

Top conviction:
- {top_names}

Action focus:
- {actions}

Main risk:
- Higher USD, yield volatility, weaker breadth.
- USD mạnh hơn, biến động lợi suất, độ rộng thị trường yếu.
"""


def generate_daily_note(
    api_key: str,
    model: str,
    note_date: str,
    market_df: pd.DataFrame,
    news_df: pd.DataFrame,
    allocation_text: str,
    consensus_text: str,
    user_notes: str,
) -> str:
    system_prompt = PROMPT_MASTER + """
You must write STRICTLY BILINGUAL output.

MANDATORY FORMAT:
- Every analytical point must be written in 2 lines:
  Line 1: Vietnamese
  Line 2: English
- Do NOT write English only.
- Do NOT group all Vietnamese first and all English later.
- For every section, Vietnamese must appear immediately before its English translation.

Write a DAILY INVESTMENT NOTE for CIO / PM.

Required structure:

1. Global Fixed Income
2. Vietnam Fixed Income
3. Global Equity
4. Vietnam Equity
5. Commodity / FX
6. Cross-asset insight
7. Top risks today
8. Action summary

For sections 1-5, each section must include:
- View
- What changed
- Why it matters
- Portfolio implication
- Risk
- Trigger
- Conviction
- Time horizon
"""
    user_prompt = f"""
Date: {note_date}

MARKET:
{build_market_highlights(market_df)}

NEWS:
{build_news_brief(news_df)}

ALLOCATION:
{allocation_text}

CONSENSUS:
{consensus_text}

USER NOTES:
{user_notes}
"""
    out = cached_ai_call(api_key, model, system_prompt, user_prompt, 2200)
    if out:
        return out

    return f"""DAILY NOTE / BẢN TIN NGÀY - {note_date}

1. Fixed Income toàn cầu giữ quan điểm trung lập do lợi suất còn biến động.
Global fixed income remains neutral as yields are still volatile.

2. Fixed Income Việt Nam cần theo dõi thêm thanh khoản và tỷ giá.
Vietnam fixed income needs further monitoring of liquidity and FX.

3. Equity toàn cầu thiên về chọn lọc hơn là tăng beta diện rộng.
Global equity favors selective exposure rather than broad beta expansion.

4. Equity Việt Nam phù hợp với large caps và các mã có đồng thuận cao.
Vietnam equity is better suited to large caps and names with stronger consensus.

5. Hàng hóa và FX tiếp tục là lớp tín hiệu quan trọng cho định vị ngắn hạn.
Commodities and FX remain key signals for short-term positioning.

6. Hàm ý danh mục:
Portfolio implication:
{allocation_text}

7. Đồng thuận cổ phiếu Việt Nam:
Vietnam stock consensus:
{consensus_text}
"""


def generate_ic_note(
    api_key: str,
    model: str,
    note_date: str,
    market_df: pd.DataFrame,
    news_df: pd.DataFrame,
    allocation_text: str,
    trade_ideas_text: str,
) -> str:
    system_prompt = PROMPT_MASTER + """
You must write STRICTLY BILINGUAL output.

MANDATORY FORMAT:
- Every analytical point must be written in 2 lines:
  Line 1: Vietnamese
  Line 2: English
- Do NOT write English only.
- Do NOT group all Vietnamese first and all English later.
- Vietnamese must be immediately followed by its English translation.

Write a STRICTLY BILINGUAL IC NOTE.

Required structure:
1. Base case
2. Bull case
3. Bear case
4. Portfolio implication
5. Top trade ideas
6. Key risks
7. Decision points
"""
    user_prompt = f"""
Date: {note_date}

MARKET:
{build_market_highlights(market_df)}

NEWS:
{build_news_brief(news_df)}

ALLOCATION:
{allocation_text}

TRADE IDEAS:
{trade_ideas_text}
"""
    out = cached_ai_call(api_key, model, system_prompt, user_prompt, 2200)
    if out:
        return out

    return f"""IC NOTE / GHI CHÚ IC - {note_date}

1. Kịch bản cơ sở là thị trường tiếp tục phân hóa, ưu tiên vị thế chọn lọc thay vì tăng beta mạnh.
The base case is a still-divergent market, favoring selective exposure rather than aggressive beta expansion.

2. Kịch bản tích cực là dòng tiền cải thiện và nhóm dẫn dắt duy trì sức mạnh.
The bull case is improving flows with leadership groups maintaining strength.

3. Kịch bản tiêu cực là USD mạnh hơn, lợi suất tăng và khẩu vị rủi ro suy yếu.
The bear case is a stronger USD, higher yields, and weaker risk appetite.

4. Hàm ý danh mục là giữ core exposure ở các tài sản có conviction cao và hạn chế đuổi giá.
The portfolio implication is to maintain core exposure in high-conviction assets while avoiding chasing.

5. Các ý tưởng nổi bật nên tập trung vào các mã có consensus tốt và phù hợp regime hiện tại.
Top ideas should focus on names with strong consensus and alignment with the current regime.

6. Rủi ro chính là biến động tỷ giá, lợi suất và độ rộng thị trường suy yếu.
The main risks are FX volatility, yields, and weakening market breadth.

7. Điểm quyết định là liệu động lượng thị trường có được xác nhận thêm bởi dòng tiền và leadership hay không.
The decision point is whether market momentum is further confirmed by flows and leadership.
"""


def main():
    ensure_files()
    st.set_page_config(page_title=APP_TITLE, layout="wide")

    if "user" not in st.session_state:
        st.session_state["user"] = None

    if st.session_state["user"] is None:
        st.title(APP_TITLE)
        st.subheader("Đăng nhập / Login")
        u = st.text_input("Username")
        p = st.text_input("Password", type="password")
        if st.button("Login", use_container_width=True):
            user = authenticate(u, p)
            if user:
                st.session_state["user"] = user
                st.rerun()
            else:
                st.error("Sai tài khoản hoặc mật khẩu.")
        return

    api_key = get_secret("OPENAI_API_KEY", "")
    model = st.sidebar.text_input("Model", value="gpt-4.1-mini")
    note_date = str(st.sidebar.date_input("Ngày báo cáo"))
    user_notes = st.sidebar.text_area("Ghi chú thêm", "")

    st.title(APP_TITLE)
    if st.button("Logout"):
        st.session_state["user"] = None
        st.rerun()

    if st.button("Run All", use_container_width=True):
        market_df = fetch_market_snapshot(DEFAULT_TICKERS)
        news_df = fetch_news()
        consensus_df = summarize_consensus()
        consensus_text = build_consensus_text(consensus_df)
        allocation = compute_allocation(market_df, consensus_df)
        allocation_text = format_allocation(allocation)
        trade_df = build_trade_ideas(market_df, consensus_df)
        trade_text = format_trade_ideas(trade_df)
        benchmark_df = compute_benchmark(market_df)
        heatmap_df = build_heatmap(trade_df)
        action_df = generate_action_signals(trade_df)
        memo_text = build_memo(allocation_text, trade_df, action_df, benchmark_df)
        daily_note = generate_daily_note(api_key, model, note_date, market_df, news_df, allocation_text, consensus_text, user_notes)
        ic_note = generate_ic_note(api_key, model, note_date, market_df, news_df, allocation_text, trade_text)

        st.session_state["market_df"] = market_df
        st.session_state["news_df"] = news_df
        st.session_state["consensus_df"] = consensus_df
        st.session_state["allocation_text"] = allocation_text
        st.session_state["trade_df"] = trade_df
        st.session_state["trade_text"] = trade_text
        st.session_state["benchmark_df"] = benchmark_df
        st.session_state["heatmap_df"] = heatmap_df
        st.session_state["action_df"] = action_df
        st.session_state["memo_text"] = memo_text
        st.session_state["daily_note"] = daily_note
        st.session_state["ic_note"] = ic_note

    tabs = st.tabs([
        "Dashboard", "Daily Note", "Allocation", "Trade Ideas",
        "Benchmark", "Heatmap", "Portfolio Memo", "Action Signals", "IC Note"
    ])

    with tabs[0]:
        if "market_df" in st.session_state:
            st.dataframe(st.session_state["market_df"], use_container_width=True)
        if "news_df" in st.session_state and not st.session_state["news_df"].empty:
            show = st.session_state["news_df"].copy()
            show["ArticleLink"] = show["Link"].apply(lambda x: f'<a href="{x}" target="_blank">Open</a>' if x else "")
            st.write(
                show[["Region", "Source", "Title", "AssetClass", "VNImpact", "ArticleLink"]].to_html(
                    escape=False, index=False
                ),
                unsafe_allow_html=True,
            )

    with tabs[1]:
        txt = st.session_state.get("daily_note", "")
        st.text_area("Daily Note (Việt - Anh)", txt, height=520)
        if txt:
            docx = export_docx("Daily Note", txt)
            pdf = export_pdf("Daily Note", txt)
            c1, c2 = st.columns(2)
            with c1:
                if docx:
                    st.download_button("Export Word", data=docx, file_name=f"daily_note_{note_date}.docx")
            with c2:
                if pdf:
                    st.download_button("Export PDF", data=pdf, file_name=f"daily_note_{note_date}.pdf")

    with tabs[2]:
        st.text_area("Allocation", st.session_state.get("allocation_text", ""), height=420)

    with tabs[3]:
        if "trade_df" in st.session_state:
            st.dataframe(st.session_state["trade_df"], use_container_width=True)
        st.text_area("Trade Ideas", st.session_state.get("trade_text", ""), height=420)

    with tabs[4]:
        if "benchmark_df" in st.session_state:
            b = st.session_state["benchmark_df"]
            st.dataframe(b, use_container_width=True)
            st.metric("Model Portfolio Return", f"{b.attrs.get('portfolio_return', 0.0)}%")
            st.metric("VNINDEX Return", f"{b.attrs.get('vnindex_return', 0.0)}%")
            st.metric("Excess Return", f"{b.attrs.get('excess_return', 0.0)}%")

    with tabs[5]:
        if "heatmap_df" in st.session_state:
            st.dataframe(st.session_state["heatmap_df"], use_container_width=True)

    with tabs[6]:
        memo = st.session_state.get("memo_text", "")
        st.text_area("Portfolio Memo", memo, height=260)
        if memo:
            docx = export_docx("Portfolio Memo", memo)
            pdf = export_pdf("Portfolio Memo", memo)
            c1, c2 = st.columns(2)
            with c1:
                if docx:
                    st.download_button("Export Memo Word", data=docx, file_name=f"portfolio_memo_{note_date}.docx")
            with c2:
                if pdf:
                    st.download_button("Export Memo PDF", data=pdf, file_name=f"portfolio_memo_{note_date}.pdf")

    with tabs[7]:
        if "action_df" in st.session_state:
            st.dataframe(st.session_state["action_df"], use_container_width=True)

    with tabs[8]:
        txt = st.session_state.get("ic_note", "")
        st.text_area("IC Note (Việt - Anh)", txt, height=520)
        if txt:
            docx = export_docx("IC Note", txt)
            pdf = export_pdf("IC Note", txt)
            c1, c2 = st.columns(2)
            with c1:
                if docx:
                    st.download_button("Export IC Word", data=docx, file_name=f"ic_note_{note_date}.docx")
            with c2:
                if pdf:
                    st.download_button("Export IC PDF", data=pdf, file_name=f"ic_note_{note_date}.pdf")


if __name__ == "__main__":
    main()
